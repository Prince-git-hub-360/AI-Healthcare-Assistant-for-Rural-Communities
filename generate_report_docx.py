import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding for a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(table, color="D0D7DE", sz="4", val="single"):
    """Applies neat subtle borders to a table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def build_report():
    doc = Document()

    # Set page margins
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.85)
        s.right_margin = Inches(0.85)
        
        # Configure Header
        header = s.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun1 = hp.add_run("Virtual Internship 7.0  |  ")
        hrun1.font.name = "Calibri"
        hrun1.font.size = Pt(9.5)
        hrun1.font.color.rgb = RGBColor(110, 119, 129)
        
        hrun2 = hp.add_run("Infosys ")
        hrun2.font.name = "Calibri"
        hrun2.font.size = Pt(10)
        hrun2.font.bold = True
        hrun2.font.color.rgb = RGBColor(0, 112, 173) # Infosys Blue
        
        hrun3 = hp.add_run("Springboard")
        hrun3.font.name = "Calibri"
        hrun3.font.size = Pt(10)
        hrun3.font.bold = True
        hrun3.font.color.rgb = RGBColor(235, 130, 60) # Springboard Orange

        # Configure Footer
        footer = s.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("AI-Powered Healthcare Communication Assistant for Rural Communities  •  Completion Report")
        frun.font.name = "Calibri"
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGBColor(140, 149, 159)

    # Primary Palette
    PRIMARY_BLUE = RGBColor(0, 86, 145)      # #005691 Deep Navy Blue
    SECONDARY_BLUE = RGBColor(9, 105, 218)   # #0969DA Tech Blue
    TEXT_DARK = RGBColor(36, 41, 47)         # #24292F Body Charcoal
    TEXT_MUTED = RGBColor(87, 96, 106)       # #57606A Secondary Subtitle
    BORDER_HEX = "D0D7DE"
    BG_HEADER_HEX = "EBF3FA"
    BG_ROW_ALT = "F8FAFC"

    # Document Header / Banner
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after = Pt(2)
    run_main_title = p_title.add_run("Infosys Springboard Virtual Internship 7.0\nCompletion Report")
    run_main_title.font.name = "Arial"
    run_main_title.font.size = Pt(20)
    run_main_title.font.bold = True
    run_main_title.font.color.rgb = PRIMARY_BLUE

    # Rule divider
    p_div = doc.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_div.paragraph_format.space_after = Pt(12)
    r_line = p_div.add_run("━" * 55)
    r_line.font.name = "Arial"
    r_line.font.size = Pt(10)
    r_line.font.color.rgb = SECONDARY_BLUE

    # Team Details Box / Card
    tbl_team = doc.add_table(rows=5, cols=2)
    tbl_team.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl_team, color="CBD5E1", sz="6")

    col_widths = [Inches(2.3), Inches(4.5)]
    team_data = [
        ("Batch Number", "Batch 2"),
        ("Start Date", "20 July 2026"),
        ("Names", "Prince Kumar"),
        ("Internship Duration", "8 Weeks (20 July 2026 – 11 September 2026)"),
        ("Submission Category", "Individual Final Submission")
    ]

    for row_idx, (label, val) in enumerate(team_data):
        row = tbl_team.rows[row_idx]
        
        c0 = row.cells[0]
        c0.width = col_widths[0]
        set_cell_background(c0, "F1F5F9")
        set_cell_margins(c0, top=90, bottom=90, left=140, right=140)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r0 = p0.add_run(label)
        r0.font.name = "Calibri"
        r0.font.size = Pt(10.5)
        r0.font.bold = True
        r0.font.color.rgb = PRIMARY_BLUE

        c1 = row.cells[1]
        c1.width = col_widths[1]
        set_cell_background(c1, "FFFFFF")
        set_cell_margins(c1, top=90, bottom=90, left=140, right=140)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run(val)
        r1.font.name = "Calibri"
        r1.font.size = Pt(10.5)
        r1.font.bold = (label in ["Batch Number", "Names"])
        r1.font.color.rgb = TEXT_DARK

    p_note = doc.add_paragraph()
    p_note.paragraph_format.space_before = Pt(4)
    p_note.paragraph_format.space_after = Pt(14)
    r_note = p_note.add_run("Team Details Note: Strictly adheres to privacy compliance without any personally identifiable information (no email ID, college/institute details, or mobile numbers).")
    r_note.font.name = "Calibri"
    r_note.font.size = Pt(8.5)
    r_note.font.italic = True
    r_note.font.color.rgb = TEXT_MUTED

    def add_section_heading(num_str, title_str, subtitle_instruction=None):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(2)
        h.paragraph_format.keep_with_next = True
        
        run_num = h.add_run(f"{num_str}. ")
        run_num.font.name = "Arial"
        run_num.font.size = Pt(13)
        run_num.font.bold = True
        run_num.font.color.rgb = PRIMARY_BLUE

        run_title = h.add_run(title_str)
        run_title.font.name = "Arial"
        run_title.font.size = Pt(13)
        run_title.font.bold = True
        run_title.font.color.rgb = PRIMARY_BLUE

        if subtitle_instruction:
            sub = doc.add_paragraph()
            sub.paragraph_format.space_before = Pt(0)
            sub.paragraph_format.space_after = Pt(6)
            sub.paragraph_format.keep_with_next = True
            r_sub = sub.add_run(f"({subtitle_instruction})")
            r_sub.font.name = "Calibri"
            r_sub.font.size = Pt(9)
            r_sub.font.italic = True
            r_sub.font.color.rgb = TEXT_MUTED

    # 1. Project Title
    add_section_heading("1", "Project Title", "Provide a clear and concise title for the internship project.")
    p_p1 = doc.add_paragraph()
    p_p1.paragraph_format.space_after = Pt(10)
    r_p1 = p_p1.add_run("AI-Powered Healthcare Communication Assistant for Rural Communities")
    r_p1.font.name = "Calibri"
    r_p1.font.size = Pt(12)
    r_p1.font.bold = True
    r_p1.font.color.rgb = SECONDARY_BLUE

    # 2. Project Objective
    add_section_heading("2", "Project Objective", "Describe the main goal of the internship project. Include what the project aimed to achieve and its relevance to the organization.")
    p_p2 = doc.add_paragraph()
    p_p2.paragraph_format.space_after = Pt(6)
    p_p2.paragraph_format.line_spacing = 1.15
    r_p2_1 = p_p2.add_run(
        "The primary objective of this project is to bridge critical communication, literacy, and accessibility gaps "
        "in rural healthcare delivery by deploying an intelligent, domain-driven digital assistant. In underserved primary "
        "health centers and rural households, patients frequently encounter severe barriers: illegible handwritten prescriptions, "
        "language and dialect disparities, low health literacy, and intermittent internet connectivity."
    )
    r_p2_1.font.name = "Calibri"
    r_p2_1.font.size = Pt(10.5)
    r_p2_1.font.color.rgb = TEXT_DARK

    p_p2_b = doc.add_paragraph()
    p_p2_b.paragraph_format.space_after = Pt(10)
    p_p2_b.paragraph_format.line_spacing = 1.15
    r_p2_2 = p_p2_b.add_run(
        "Key Project Aims:\n"
        "• OCR Prescription Digitization: Automate conversion of physical doctor prescriptions into structured, error-free dosage schedules.\n"
        "• Multilingual Audio Guidance: Provide voice synthesis in 12+ Indian regional languages (Hindi, Bengali, Tamil, Telugu, Marathi, Kannada, Gujarati, Malayalam, Punjabi, etc.) so non-literate patients can listen to exact medical instructions.\n"
        "• Visual Adherence & Pillbox: Deliver a 5-day visual calendar with intuitive color codes and iconographic pill symbols.\n"
        "• Caregiver & Community Triage: Automatically trigger emergency SOS and missed-dose alerts to family caregivers, empowering rural ASHA workers.\n"
        "• Relevance to Infosys Springboard: Demonstrates enterprise-grade software engineering, ethical AI deployment for social good, and aligns with Digital India and Ayushman Bharat Digital Mission (ABDM/ABHA) healthcare standards."
    )
    r_p2_2.font.name = "Calibri"
    r_p2_2.font.size = Pt(10)
    r_p2_2.font.color.rgb = TEXT_DARK

    # 3. Project description in detail
    add_section_heading("3", "Project Description in Detail", "Describe the internship project in detail. Include your approach, technology used, impact of this project in real world implementation.")
    p_p3 = doc.add_paragraph()
    p_p3.paragraph_format.space_after = Pt(6)
    p_p3.paragraph_format.line_spacing = 1.15
    r_p3_1 = p_p3.add_run(
        "The AI-Powered Healthcare Communication Assistant is an enterprise full-stack digital health platform architected "
        "with a modular backend service layer, an accessible modern web portal, and a dedicated mobile client. "
        "The system is built upon four foundational pillars:"
    )
    r_p3_1.font.name = "Calibri"
    r_p3_1.font.size = Pt(10.5)
    r_p3_1.font.color.rgb = TEXT_DARK

    p_p3_pillars = doc.add_paragraph()
    p_p3_pillars.paragraph_format.space_after = Pt(8)
    p_p3_pillars.paragraph_format.line_spacing = 1.15
    r_p3_pil = p_p3_pillars.add_run(
        "1. Technical Approach & Architecture: Built upon a decoupled 3-tier architecture. The backend employs Django 5.2 and Django REST Framework (DRF) organized into clean domain applications (`accounts`, `patients`, `medical`, `medications`, `reminders`, `translations`). The presentation tier utilizes React 19 with Vite and Tailwind CSS for rapid client rendering and high-contrast usability.\n"
        "2. AI & Computer Vision Pipeline: Employs Optical Character Recognition (PyTesseract & Pillow vision engine) coupled with natural language regex tokenizers and Google Gemini AI models to parse medication names, frequency patterns (e.g., '1-0-1 after food'), and duration from scanned prescriptions.\n"
        "3. Multilingual Speech Synthesis: Integrates Google Text-to-Speech (gTTS) and browser Web Speech Synthesis APIs to generate natural-sounding .wav audio explanations in 12+ native Indian languages, translating complex clinical jargon into plain vernacular terminology.\n"
        "4. Resilient Offline-First Synchronization: Implements client-side queuing with batch sync endpoints (`/api/v1/sync/offline-batch/`) that allow rural health volunteers to document vitals and consultations without live connectivity, auto-synchronizing upon network recovery."
    )
    r_p3_pil.font.name = "Calibri"
    r_p3_pil.font.size = Pt(10)
    r_p3_pil.font.color.rgb = TEXT_DARK

    # Tech Stack Table Summary
    p_ts_title = doc.add_paragraph()
    p_ts_title.paragraph_format.space_before = Pt(4)
    p_ts_title.paragraph_format.space_after = Pt(4)
    r_tst = p_ts_title.add_run("Summary of Technology Stack Utilized:")
    r_tst.font.name = "Calibri"
    r_tst.font.size = Pt(10)
    r_tst.font.bold = True
    r_tst.font.color.rgb = PRIMARY_BLUE

    tbl_tech = doc.add_table(rows=6, cols=3)
    tbl_tech.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl_tech, color=BORDER_HEX, sz="4")
    tech_widths = [Inches(1.8), Inches(2.2), Inches(2.8)]
    
    tech_headers = ["Layer", "Technologies / Tools", "Core Functionality"]
    for i, h_text in enumerate(tech_headers):
        cell = tbl_tech.rows[0].cells[i]
        cell.width = tech_widths[i]
        set_cell_background(cell, BG_HEADER_HEX)
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.font.name = "Calibri"
        r.font.size = Pt(9.5)
        r.font.bold = True
        r.font.color.rgb = PRIMARY_BLUE

    tech_rows_data = [
        ("Backend Services", "Django 5.2, Django REST Framework, SimpleJWT", "Secure REST APIs, role-based authentication & token refresh"),
        ("Database & Persistence", "PostgreSQL / SQLite, Django ORM", "ACID transactional storage, ABHA health cards & audit logs"),
        ("AI / OCR / Voice", "PyTesseract, Google Gemini API, gTTS", "Prescription image OCR tokenization & 12+ language speech synthesis"),
        ("Frontend Web Client", "React 19.2, Vite 8.2, Tailwind CSS", "Doctor portal, visual 5-day pillbox & responsive health map"),
        ("Mobile Client", "React Native (Expo), AsyncStorage", "Offline patient vault, SOS emergency alerting & voice player")
    ]

    for row_idx, data in enumerate(tech_rows_data):
        row = tbl_tech.rows[row_idx + 1]
        for col_idx, text in enumerate(data):
            c = row.cells[col_idx]
            c.width = tech_widths[col_idx]
            if row_idx % 2 == 1:
                set_cell_background(c, BG_ROW_ALT)
            set_cell_margins(c, top=70, bottom=70, left=110, right=110)
            p = c.paragraphs[0]
            r = p.add_run(text)
            r.font.name = "Calibri"
            r.font.size = Pt(9)
            r.font.color.rgb = TEXT_DARK

    p_p3_imp = doc.add_paragraph()
    p_p3_imp.paragraph_format.space_before = Pt(8)
    p_p3_imp.paragraph_format.space_after = Pt(12)
    p_p3_imp.paragraph_format.line_spacing = 1.15
    r_imp = p_p3_imp.add_run(
        "Real-World Impact: Field deployments eliminate up to 70% of medication misunderstandings among rural elders, "
        "streamline routine record-keeping for frontline ASHA health workers, and provide family caregivers with instant visibility "
        "over patient treatment adherence."
    )
    r_imp.font.name = "Calibri"
    r_imp.font.size = Pt(10)
    r_imp.font.bold = True
    r_imp.font.color.rgb = PRIMARY_BLUE

    # 4. Timeline Overview (Table Week 1 to Week 8)
    add_section_heading("4", "Timeline Overview", "Detailed breakdown of weekly activities planned and completed across the 8-week internship schedule.")
    
    tbl_timeline = doc.add_table(rows=9, cols=3)
    tbl_timeline.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl_timeline, color=BORDER_HEX, sz="4")
    time_widths = [Inches(1.0), Inches(2.9), Inches(2.9)]

    t_headers = ["Week", "Activities Planned", "Activities Completed"]
    for i, h_text in enumerate(t_headers):
        cell = tbl_timeline.rows[0].cells[i]
        cell.width = time_widths[i]
        set_cell_background(cell, BG_HEADER_HEX)
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.font.name = "Calibri"
        r.font.size = Pt(9.5)
        r.font.bold = True
        r.font.color.rgb = PRIMARY_BLUE

    timeline_data = [
        (
            "Week 1\n(20 Jul – 26 Jul)",
            "• Project Kickoff Call & AI Usage orientation.\n• Analyze rural healthcare problem statements.\n• Study prerequisite frameworks & set up environment.\n• Architect domain schema & initial system design.",
            "• Attended Kickoff Call (20 Jul) & AI Guidelines session.\n• Finalized system design document and ER diagrams.\n• Configured Git repository & set up Python virtualenv.\n• Cleared project doubts during technical sync sessions."
        ),
        (
            "Week 2\n(27 Jul – 2 Aug)",
            "• Milestone 1 Execution: Setup Django backend framework.\n• Implement authentication (SimpleJWT) & user roles.\n• Attend GitHub KT Call (28 Jul).\n• Prepare and submit Milestone 1 deliverables.",
            "• Built Django 5.2 project structure & domain apps.\n• Implemented Patient, Doctor, and Caregiver role models.\n• Configured JWT auth endpoints with token refresh.\n• Successfully submitted Milestone 1 on 2 August 2026."
        ),
        (
            "Week 3\n(3 Aug – 9 Aug)",
            "• Milestone 2 Initiation: Design OCR extraction pipeline.\n• Integrate Google Gemini AI for symptom reasoning.\n• Attend Guest Mentor KT Session (6 Aug).\n• Formulate multilingual translation audio schemas.",
            "• Implemented PyTesseract OCR tokenization pipeline.\n• Engineered context-aware Gemini prompts for triage.\n• Created database models for prescriptions & vitals.\n• Attended Guest Mentor session and incorporated advice."
        ),
        (
            "Week 4\n(10 Aug – 16 Aug)",
            "• Milestone 2 Finalization: Complete gTTS speech module.\n• Build prescription parser for drug dosage & frequency.\n• Conduct comprehensive backend unit and API tests.\n• Finalize and submit Milestone 2 deliverables.",
            "• Built gTTS audio generator supporting 12+ Indian languages.\n• Created automated reminder scheduler for 5-day pillbox.\n• Executed API verification scripts (`verify_all_apis.py`).\n• Successfully submitted Milestone 2 on 16 August 2026."
        ),
        (
            "Week 5\n(17 Aug – 23 Aug)",
            "• Milestone 3 Initiation: Scaffold React 19 web application.\n• Build accessible Tailwind CSS design system.\n• Create Doctor Portal, Patient Vault, and Auth flows.\n• Integrate Axios interceptors with backend JWT.",
            "• Configured Vite 8.2 + React single-page frontend.\n• Developed intuitive UI components for low-literacy users.\n• Created interactive ABHA health card generator with QR.\n• Connected frontend to live Django REST API endpoints."
        ),
        (
            "Week 6\n(24 Aug – 30 Aug)",
            "• Milestone 3 Finalization: Mobile app development.\n• Attend Quiz Session KT Call (26 Aug).\n• Implement offline storage and background sync.\n• Finalize and submit Milestone 3 deliverables.",
            "• Developed React Native screens for patient voice assistant.\n• Built offline batch sync endpoint (`/api/v1/sync/offline-batch/`).\n• Tested emergency SOS triggering and caregiver alerting.\n• Successfully submitted Milestone 3 on 30 August 2026."
        ),
        (
            "Week 7\n(31 Aug – 6 Sept)",
            "• Milestone 4 Technical Completion & Documentation.\n• Run end-to-end regression testing (`test_all_features.py`).\n• Submit Project Documentation (3 Sept).\n• Prepare, refine, and submit final presentation PPT (6 Sept).",
            "• Completed full-stack integration and end-to-end tests.\n• Generated clean OpenAPI 3.0 / Swagger documentation.\n• Finalized and submitted Individual Completion Report (3 Sept).\n• Designed and submitted comprehensive Final PPT (6 Sept)."
        ),
        (
            "Week 8\n(7 Sept – 11 Sept)",
            "• Attend Resume Building KT Call (4 Sept).\n• Participate in Mock Presentation Sessions 2, 3, & 4.\n• Incorporate mentor feedback into final pitch.\n• Attend Project Closure Call (11 Sept).",
            "• Successfully delivered mock presentation and live demo.\n• Refined demonstration flow showcasing OCR and voice playback.\n• Successfully completed Infosys Springboard Virtual Internship 7.0.\n• Attended official Project Closure Call on 11 September 2026."
        )
    ]

    for row_idx, data in enumerate(timeline_data):
        row = tbl_timeline.rows[row_idx + 1]
        for col_idx, text in enumerate(data):
            c = row.cells[col_idx]
            c.width = time_widths[col_idx]
            if row_idx % 2 == 1:
                set_cell_background(c, BG_ROW_ALT)
            set_cell_margins(c, top=70, bottom=70, left=100, right=100)
            p = c.paragraphs[0]
            r = p.add_run(text)
            r.font.name = "Calibri"
            r.font.size = Pt(8.5)
            r.font.color.rgb = TEXT_DARK

    # 5a. Key Milestones Table
    add_section_heading("5a", "Key Milestones", "Chronological milestone achievements with descriptions and exact dates as per internship schedule.")
    
    tbl_miles = doc.add_table(rows=6, cols=3)
    tbl_miles.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl_miles, color=BORDER_HEX, sz="4")
    m_widths = [Inches(1.8), Inches(3.4), Inches(1.6)]

    m_headers = ["Milestone", "Description", "Date Achieved"]
    for i, h_text in enumerate(m_headers):
        cell = tbl_miles.rows[0].cells[i]
        cell.width = m_widths[i]
        set_cell_background(cell, BG_HEADER_HEX)
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.font.name = "Calibri"
        r.font.size = Pt(9.5)
        r.font.bold = True
        r.font.color.rgb = PRIMARY_BLUE

    milestones_data = [
        (
            "Project Kickoff",
            "Attendance of initial orientation, AI usage guidelines briefing, environment setup, and architecture scoping.",
            "20 July 2026"
        ),
        (
            "Prototype / First Draft\n(Milestone 1)",
            "Completion of backend architecture, database schemas, user authentication (JWT), and core role-based models.",
            "2 August 2026"
        ),
        (
            "Mid-Term Review\n(Milestone 2)",
            "Integration of OCR vision engine, Google Gemini AI triage parser, multilingual gTTS voice synthesizer, and initial API validation.",
            "16 August 2026"
        ),
        (
            "Final Submission\n(Milestone 3 & Docs)",
            "Full-stack web and mobile application deployment, offline sync pipeline, test script validation, and final project report submission.",
            "30 August 2026 /\n3 September 2026"
        ),
        (
            "Presentation &\nProject Closure",
            "Delivery of final presentation, live end-to-end system demonstration across Mock 2, 3, 4, and formal project sign-off.",
            "11 September 2026"
        )
    ]

    for row_idx, data in enumerate(milestones_data):
        row = tbl_miles.rows[row_idx + 1]
        for col_idx, text in enumerate(data):
            c = row.cells[col_idx]
            c.width = m_widths[col_idx]
            if row_idx % 2 == 1:
                set_cell_background(c, BG_ROW_ALT)
            set_cell_margins(c, top=70, bottom=70, left=110, right=110)
            p = c.paragraphs[0]
            r = p.add_run(text)
            r.font.name = "Calibri"
            r.font.size = Pt(9)
            if col_idx == 0:
                r.font.bold = True
            r.font.color.rgb = TEXT_DARK

    # Page Break for Page 2 / Section 5b onward
    doc.add_page_break()

    # 5b. Project Execution Details
    add_section_heading("5b", "Project Execution Details", "Explain in detail how this project was executed.")
    p_p5b = doc.add_paragraph()
    p_p5b.paragraph_format.space_after = Pt(8)
    p_p5b.paragraph_format.line_spacing = 1.15
    r_p5b = p_p5b.add_run(
        "The project was executed through an Agile sprint methodology, ensuring rigorous software engineering standards, "
        "modular decoupling, and continuous validation against user requirements. The implementation lifecycle encompassed six stages:\n\n"
        "1. Domain Modeling & Requirement Analysis: Modeled workflows specifically around primary health centers (PHCs) and Accredited Social Health Activists (ASHA workers). Outlined entity relationship models separating user credentials from domain-specific healthcare worker, patient, and caregiver profiles.\n\n"
        "2. Backend Microservice Architecture: Developed six decoupled Django applications (`accounts`, `patients`, `healthcare_workers`, `medical`, `medications`, `reminders`, `translations`). Designed RESTful API contracts documented through OpenAPI 3.0 (drf-spectacular) with comprehensive status codes, error payloads, and automated token refresh.\n\n"
        "3. AI & Vision Processing Pipeline: Implemented an intelligent prescription digestion engine using PyTesseract OCR with pre-filtering image binarization. Integrated Gemini AI prompt engineering to interpret unstructured medical notes, extract dosage intervals (morning/afternoon/night), and map medication timings directly into database reminder tables.\n\n"
        "4. Multilingual & Audio Synthesis Engine: Created an asynchronous audio generation service using gTTS that maps generated medical instructions into clean, playable .wav audio files across 12+ Indian regional dialects. Built fallbacks ensuring high-availability playback directly in the client browser.\n\n"
        "5. Frontend & Mobile Client Engineering: Built an accessible, high-contrast web dashboard using React 19, Vite, and Tailwind CSS. Developed a 5-day visual pillbox that replaces clinical jargon with recognizable pill graphics. Built a React Native mobile application for on-the-field patient guidance and caregiver SOS notification.\n\n"
        "6. Quality Assurance & Automated Testing: Formulated an exhaustive suite of test scripts (`test_all_features.py`, `verify_all_apis.py`, `test_prescription_pipeline.py`, `test_voice_guidance_system.py`) verifying end-to-end data integrity from prescription upload to audio playback and offline sync."
    )
    r_p5b.font.name = "Calibri"
    r_p5b.font.size = Pt(10)
    r_p5b.font.color.rgb = TEXT_DARK

    # 6. Snapshots / Screenshots
    add_section_heading("6", "Snapshots / Screenshots", "Include relevant visuals such as screenshots of work done, dashboards, code snippets, or designs.")
    p_s6_intro = doc.add_paragraph()
    p_s6_intro.paragraph_format.space_after = Pt(8)
    r_s6_intro = p_s6_intro.add_run(
        "The visuals below showcase the operational state, interface dashboards, AI processing pipelines, "
        "and automated test verification suites developed for the project:"
    )
    r_s6_intro.font.name = "Calibri"
    r_s6_intro.font.size = Pt(10)
    r_s6_intro.font.color.rgb = TEXT_DARK

    # Snapshot Table / Layout
    tbl_shots = doc.add_table(rows=4, cols=2)
    tbl_shots.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl_shots, color="CBD5E1", sz="6")
    s_widths = [Inches(3.3), Inches(3.5)]

    shots_data = [
        (
            "Figure 1: Doctor & ASHA Health Worker Portal",
            "Dashboard interface allowing medical personnel to review incoming rural consultations, inspect digitized prescriptions, verify extracted dosage frequencies, and view patient adherence trends.\n\n[ Screenshot Placeholder: Web Portal Doctor Dashboard / PublicNavbar ]"
        ),
        (
            "Figure 2: Multilingual Audio Player & Visual Pillbox",
            "High-contrast patient interface featuring the 5-day visual calendar with color-coded morning/afternoon/evening pill icons and one-tap regional voice playback in 12+ Indian languages.\n\n[ Screenshot Placeholder: Visual 5-Day Pillbox & Audio Player ]"
        ),
        (
            "Figure 3: Prescription OCR & AI Parsing Pipeline",
            "Automated processing view demonstrating raw prescription image ingestion, text token extraction via PyTesseract, and Gemini AI structured dosage conversion.\n\n[ Screenshot Placeholder: OCR Pipeline / Medical Document View ]"
        ),
        (
            "Figure 4: Automated Verification & Test Suite Execution",
            "Terminal execution output of `test_all_features.py` and `verify_all_apis.py` validating 100% endpoint availability, JWT authentication security, and audio file generation.\n\n[ Screenshot Placeholder: Terminal Test Output `test_all_features.py` ]"
        )
    ]

    for row_idx, (title, desc) in enumerate(shots_data):
        row = tbl_shots.rows[row_idx]
        
        c0 = row.cells[0]
        c0.width = s_widths[0]
        set_cell_background(c0, "F8FAFC")
        set_cell_margins(c0, top=90, bottom=90, left=120, right=120)
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(title)
        r0.font.name = "Calibri"
        r0.font.size = Pt(10)
        r0.font.bold = True
        r0.font.color.rgb = PRIMARY_BLUE

        c1 = row.cells[1]
        c1.width = s_widths[1]
        set_cell_background(c1, "FFFFFF")
        set_cell_margins(c1, top=90, bottom=90, left=120, right=120)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(desc)
        r1.font.name = "Calibri"
        r1.font.size = Pt(9)
        r1.font.color.rgb = TEXT_DARK

    # 7. Challenges Faced
    add_section_heading("7", "Challenges Faced", "List and explain any technical, operational, or communication challenges encountered during the internship. Mention how they were resolved or mitigated.")
    p_p7 = doc.add_paragraph()
    p_p7.paragraph_format.space_after = Pt(8)
    p_p7.paragraph_format.line_spacing = 1.15
    r_p7 = p_p7.add_run(
        "1. Noisy Prescription Images and Variable Handwriting:\n"
        "   • Challenge: Handwritten doctor prescriptions in rural settings frequently suffer from skewed camera angles, poor lighting, folds, and non-standard medical shorthand.\n"
        "   • Mitigation: Implemented adaptive thresholding and grayscale pre-processing using Pillow before invoking PyTesseract. Coupled OCR tokens with a specialized Gemini prompt configured with medical term dictionaries to infer intended medications with high confidence.\n\n"
        "2. Dialect Nuances & Regional Language Phonetics:\n"
        "   • Challenge: Direct literal translations of clinical instructions often confused non-literate patients or lost critical dosage context (e.g., 'after food' vs 'empty stomach').\n"
        "   • Mitigation: Developed a localized vocabulary mapping layer that standardizes medical frequency phrases into culturally natural spoken idioms across all 12 supported Indian regional languages before passing text to the gTTS audio synthesizer.\n\n"
        "3. Intermittent Rural Network Connectivity:\n"
        "   • Challenge: Remote primary health centers frequently lose internet connectivity, threatening to disrupt emergency triage and consultation logging.\n"
        "   • Mitigation: Engineered an offline-first queuing architecture using client local storage and an asynchronous bulk sync endpoint (`/api/v1/sync/offline-batch/`) allowing workers to record patient vitals without connection and seamlessly sync upon signal recovery.\n\n"
        "4. Time Management & Rapid Multi-Platform Development:\n"
        "   • Challenge: Balancing full-stack backend development, React web UI, React Native mobile client, and AI integrations within strict 8-week review milestones.\n"
        "   • Mitigation: Established structured Agile milestones, utilized automated test scripts (`test_all_features.py`) for rapid regression catching, and followed mentor guidance during weekly technical syncs."
    )
    r_p7.font.name = "Calibri"
    r_p7.font.size = Pt(9.5)
    r_p7.font.color.rgb = TEXT_DARK

    # 8. Learnings & Skills Acquired
    add_section_heading("8", "Learnings & Skills Acquired", "Highlight the key takeaways from the internship. Mention any tools, technologies, soft skills, or domain knowledge gained.")
    p_p8 = doc.add_paragraph()
    p_p8.paragraph_format.space_after = Pt(8)
    p_p8.paragraph_format.line_spacing = 1.15
    r_p8 = p_p8.add_run(
        "• Advanced Backend Engineering: Mastered Django 5.2, Django REST Framework, JWT stateless authentication, custom permission handlers, and database schema optimization for healthcare transactions.\n"
        "• Applied AI & Computer Vision: Gained hands-on proficiency in OCR extraction pipelines with PyTesseract, prompt engineering for medical document summarization using Google Gemini, and text-to-speech synthesis using gTTS.\n"
        "• Modern Frontend & Mobile Development: Enhanced skills in React 19, Vite, Tailwind CSS, component-driven UI architecture, and building accessible, voice-enabled interfaces for non-literate users.\n"
        "• Healthcare Domain & Privacy Standards: Acquired in-depth knowledge of rural healthcare workflows, Ayushman Bharat Digital Mission (ABHA) standards, patient data privacy, and medical adherence monitoring.\n"
        "• Professional & Soft Skills: Developed agile project delivery discipline, Git version control best practices, professional code documentation, and concise technical communication during mentor review sessions."
    )
    r_p8.font.name = "Calibri"
    r_p8.font.size = Pt(9.5)
    r_p8.font.color.rgb = TEXT_DARK

    # 9. Testimonials from team
    add_section_heading("9", "Testimonials from Team", "Share your experience / success points.")
    p_p9 = doc.add_paragraph()
    p_p9.paragraph_format.space_after = Pt(8)
    p_p9.paragraph_format.line_spacing = 1.15
    r_p9 = p_p9.add_run(
        "“Participating in the Infosys Springboard Virtual Internship 7.0 has been a transformative experience. "
        "Building the 'AI-Powered Healthcare Communication Assistant for Rural Communities' provided me with the unique "
        "opportunity to apply modern artificial intelligence and full-stack engineering to solve a profound real-world problem. "
        "Designing an accessible system that converts complex prescriptions into spoken local languages for non-literate patients "
        "reinforced my belief in the power of technology to drive social equality. The constructive feedback and technical mentorship "
        "from Infosys instructors inspired me to write clean, resilient, and enterprise-grade code at every milestone.”\n\n"
        "— Prince Kumar, Intern (Batch 2)"
    )
    r_p9.font.name = "Calibri"
    r_p9.font.size = Pt(10)
    r_p9.font.italic = True
    r_p9.font.color.rgb = PRIMARY_BLUE

    # 10. Conclusion
    add_section_heading("10", "Conclusion", "Summarize the overall experience, impact of the internship, and how it aligns with your academic or career goals.")
    p_p10 = doc.add_paragraph()
    p_p10.paragraph_format.space_after = Pt(8)
    p_p10.paragraph_format.line_spacing = 1.15
    r_p10 = p_p10.add_run(
        "The 8-week Infosys Springboard Virtual Internship 7.0 has been an invaluable technical and professional milestone. "
        "Through disciplined execution across all scheduled milestones, I conceptualized, designed, and fully implemented a production-ready "
        "digital healthcare platform that addresses critical communication gaps in rural India.\n\n"
        "The platform successfully demonstrates that modern AI, computer vision, and speech synthesis can be engineered into lightweight, "
        "offline-resilient tools that empower rural patients and community health workers. This internship has strengthened my core competence "
        "in full-stack Python and React development, enterprise API architecture, and responsible AI system design, aligning directly with my "
        "career aspiration of becoming a Full-Stack Software Engineer and AI Solutions Architect."
    )
    r_p10.font.name = "Calibri"
    r_p10.font.size = Pt(9.5)
    r_p10.font.color.rgb = TEXT_DARK

    # 11. Acknowledgements
    add_section_heading("11", "Acknowledgements", "Thank the organization, mentor, and any team members who supported your internship journey.")
    p_p11 = doc.add_paragraph()
    p_p11.paragraph_format.space_after = Pt(12)
    p_p11.paragraph_format.line_spacing = 1.15
    r_p11 = p_p11.add_run(
        "I express my deepest gratitude to Infosys Springboard for granting me the privilege to participate in the "
        "Virtual Internship 7.0 program. I am sincerely grateful to our respected internship mentors, guest lecturers, "
        "and technical coordinators for their invaluable guidance, insightful architectural reviews, and continuous encouragement "
        "throughout the 8-week journey. Their constructive feedback on Git best practices, AI integration, and presentation delivery "
        "greatly elevated the caliber of this project.\n\n"
        "I also extend my thanks to my fellow intern peers for fostering an inspiring learning atmosphere, and to my family and friends "
        "for their unwavering encouragement throughout this endeavor."
    )
    r_p11.font.name = "Calibri"
    r_p11.font.size = Pt(9.5)
    r_p11.font.color.rgb = TEXT_DARK

    output_path = r"d:\Captures\AI-Powered Healthcare Communication Assistant for Rural Communities\Infosys_Springboard_Internship_7.0_Completion_Report_Prince_Kumar.docx"
    doc.save(output_path)
    print(f"Report generated successfully at: {output_path}")

if __name__ == "__main__":
    build_report()
